from groq import Groq
import json
import os
from datetime import datetime
from docx import Document

TEMPLATES = {
    "standard": {
        "name": "Standard Transcript",
        "description": "Speaker labels, timestamps, paragraph breaks",
        "system_prompt": """Format this as a standard transcript with:
- Clear speaker labels (use custom names if provided)
- Timestamps at the start of each speaker's turn
- Natural paragraph breaks for readability
- Preserve the conversational flow"""
    },

    "qa": {
        "name": "Q&A Format",
        "description": "Question/Answer structure",
        "system_prompt": """Format this as a Q&A transcript with:
- Clear Q: and A: labels for questions and answers
- Use speaker names if provided to identify questioner/responder
- Group related exchanges together
- Remove timestamps unless specifically requested"""
    },

    "clean": {
        "name": "Clean Text",
        "description": "No timestamps or speaker labels, flowing paragraphs",
        "system_prompt": """Format this as clean flowing text with:
- NO speaker labels or timestamps
- Smooth transitions between speakers
- Natural paragraph structure
- Read like a written article or story"""
    },

    "captions": {
        "name": "Timestamped Captions",
        "description": "SRT/VTT caption style",
        "system_prompt": """Format this as caption-style segments with:
- Precise timestamps for each caption block
- Short segments (2-3 lines max per segment)
- Proper line breaks for readability on screen
- Format: [HH:MM:SS] Caption text"""
    }
}

def get_template_prompt(template_name):
    """Get the system prompt for a template."""
    template = TEMPLATES.get(template_name, TEMPLATES["standard"])
    return template["system_prompt"]

def build_formatting_instructions(formatting_options):
    """
    Build formatting instructions from options dict.

    Args:
        formatting_options: Dict with boolean flags for each option

    Returns:
        String with formatting instructions
    """
    instructions = []

    if formatting_options.get("filter_swears"):
        instructions.append("- Replace swear words and profanity with appropriate alternatives")

    if formatting_options.get("remove_fillers"):
        instructions.append("- Remove filler words (um, uh, like, you know, sort of)")

    if formatting_options.get("improve_punctuation"):
        instructions.append("- Add proper punctuation and capitalization")

    if formatting_options.get("clean_grammar"):
        instructions.append("- Fix grammatical errors while preserving meaning")

    if formatting_options.get("add_timestamps"):
        instructions.append("- Include timestamps from the original segments")

    if formatting_options.get("formal_tone"):
        instructions.append("- Convert to formal/technical tone where appropriate")
    else:
        instructions.append("- Preserve the original conversational tone")

    return "\n".join(instructions) if instructions else "- Preserve text as-is with minimal changes"

def build_structured_transcript(segments, formatting_options=None):
    """Build transcript with speaker labels and timestamps from segments."""
    lines = []

    for segment in segments:
        speaker = segment.get("speaker", "Speaker")
        text = segment.get("text", "")
        start = segment.get("start", 0)

        # Format timestamp
        minutes = int(start // 60)
        seconds = int(start % 60)
        timestamp = f"[{minutes:02d}:{seconds:02d}]"

        # Build line with speaker and timestamp
        if formatting_options and formatting_options.get("add_timestamps"):
            lines.append(f"{timestamp} {speaker}: {text}")
        else:
            lines.append(f"{speaker}: {text}")

    return "\n".join(lines)

def format_transcript(json_filepath, api_key, speaker_mapping=None,
                     formatting_options=None, template="standard",
                     progress_callback=None):
    """
    Format transcript using Groq API.

    Args:
        json_filepath: Path to transcript JSON file
        api_key: Groq API key
        speaker_mapping: Dict mapping speaker labels to custom names (e.g., {"SPEAKER_00": "John"})
        formatting_options: Dict of formatting options (checkboxes)
        template: Template name to use
        progress_callback: Optional callback for progress updates

    Returns:
        Formatted transcript text
    """
    if progress_callback:
        progress_callback("Loading transcript...")

    # Load transcript
    with open(json_filepath, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Build transcript with speaker structure if segments available
    segments = data.get("segments", [])

    if segments and template in ["standard", "qa", "captions"]:
        # Use structured segments for templates that need speaker info
        transcript_text = build_structured_transcript(segments, formatting_options)
    else:
        # Use simple full transcript for clean template
        transcript_text = data["full_transcript"]

    if progress_callback:
        progress_callback("Sending to Groq API...")

    # Initialize Groq client
    client = Groq(api_key=api_key)

    # Get template prompt
    template_prompt = get_template_prompt(template)

    # Build formatting instructions
    format_instructions = build_formatting_instructions(formatting_options or {})

    # Build speaker replacement instructions
    speaker_instructions = ""
    if speaker_mapping:
        replacements = [f"'{old}' should be '{new}'"
                       for old, new in speaker_mapping.items() if new.strip()]
        if replacements:
            speaker_instructions = "\nSpeaker name replacements:\n" + "\n".join(replacements) + "\n"

    # Format using Groq
    prompt = f"""{template_prompt}

Additional formatting instructions:
{format_instructions}

{speaker_instructions}

Transcript:
{transcript_text}"""

    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        model="llama-3.3-70b-versatile",
        temperature=0.3,
    )

    formatted_text = chat_completion.choices[0].message.content

    if progress_callback:
        progress_callback("Formatting complete!")

    return formatted_text

def save_to_docx(formatted_text, project_name, progress_callback=None):
    """
    Save formatted transcript to Word document.

    Args:
        formatted_text: The formatted transcript text
        project_name: Name of the project
        progress_callback: Optional callback for progress updates

    Returns:
        Path to saved .docx file
    """
    if progress_callback:
        progress_callback("Creating Word document...")

    # Create Word document
    doc = Document()
    doc.add_heading(f"Transcript: {project_name}", 0)
    doc.add_paragraph(formatted_text)

    # Save to completed folder
    os.makedirs("completed", exist_ok=True)
    timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_filename = f"completed/{project_name}_{timestamp_str}.docx"

    doc.save(docx_filename)

    if progress_callback:
        progress_callback("Document saved!")

    return docx_filename
